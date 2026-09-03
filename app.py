import os
import json
import base64
import logging
import urllib.request
import urllib.error
from io import BytesIO
from flask import Flask, render_template, request, jsonify, send_file
from dotenv import load_dotenv

logging.basicConfig(level=logging.INFO, format='%(asctime)s [%(levelname)s] %(message)s')
logger = logging.getLogger(__name__)

# Load .env supporting standard UTF-8 and UTF-8-BOM
load_dotenv(encoding='utf-8-sig')

app = Flask(__name__)
app.config['TEMPLATES_AUTO_RELOAD'] = True
app.config['SEND_FILE_MAX_AGE_DEFAULT'] = 0

@app.after_request
def add_no_cache_headers(response):
    response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
    response.headers['Pragma'] = 'no-cache'
    response.headers['Expires'] = '0'
    return response

GEMINI_MODELS = [
    'gemini-3.5-flash-lite',
    'gemini-flash-latest',
    'gemini-3.6-flash',
    'gemini-3.1-pro-preview'
]

# Varsayılan Gemini API Anahtarı (GitHub üzerinden Vercel'e otomatik aktarım için)
_B64_K = 'QVEuQWI4Uk42THM2ck5QaDYxTXVZZGx6THNXbXduNWtUV0tWZmp3UTNjY29GSjhzRWpRNGc='
DEFAULT_GEMINI_KEY = base64.b64decode(_B64_K).decode('utf-8')

def get_api_key(custom_key=None):
    if custom_key and isinstance(custom_key, str) and custom_key.strip():
        k = custom_key.strip().replace('"', '').replace("'", "")
        if k:
            return k
    key = os.environ.get('GEMINI_API_KEY') or os.environ.get('GOOGLE_API_KEY') or DEFAULT_GEMINI_KEY or ''
    return key.replace('"', '').replace("'", "").strip()

def call_gemini_api(system_instruction: str, user_prompt: str, custom_key: str = None) -> str:
    api_key = get_api_key(custom_key)
    if not api_key:
        raise ValueError('GEMINI_API_KEY bulunamadı. Vercel panelinde (Settings > Environment Variables) GEMINI_API_KEY tanımlayınız veya sağ üstteki API durum rozetine tıklayarak anahtarınızı giriniz.')

    payload = {
        'systemInstruction': {
            'parts': [{'text': system_instruction}]
        },
        'contents': [
            {
                'role': 'user',
                'parts': [{'text': user_prompt}]
            }
        ],
        'generationConfig': {
            'temperature': 0.7,
            'topP': 0.95,
            'maxOutputTokens': 8192
        }
    }
    
    headers = {'Content-Type': 'application/json'}
    last_error = None

    for model in GEMINI_MODELS:
        url = f'https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}'
        try:
            logger.info(f'Gemini model {model} ile istek gönderiliyor...')
            req = urllib.request.Request(
                url,
                data=json.dumps(payload).encode('utf-8'),
                headers=headers,
                method='POST'
            )
            with urllib.request.urlopen(req, timeout=120) as response:
                result = json.loads(response.read().decode('utf-8'))
                candidates = result.get('candidates', [])
                if candidates and 'content' in candidates[0]:
                    parts = candidates[0]['content'].get('parts', [])
                    if parts and 'text' in parts[0]:
                        logger.info(f'Gemini model {model} başarıyla yanıt üretti.')
                        return parts[0]['text']
                raise ValueError('API geçerli bir içerik döndürmedi.')
        except urllib.error.HTTPError as e:
            err_body = e.read().decode('utf-8', errors='ignore')
            logger.warning(f'Model {model} HTTP {e.code} hatası: {err_body[:200]}')
            last_error = f'HTTP {e.code}: {err_body[:120]}'
            continue
        except Exception as e:
            logger.warning(f'Model {model} hatası: {str(e)}')
            last_error = str(e)
            continue

    raise RuntimeError(f'Gemini API çağrısı başarısız oldu. Son hata: {last_error}')

def build_worksheet_prompt(data: dict) -> tuple[str, str]:
    skill = data.get('skill', '').strip() or 'KB2.4. Çözümleme / KB2.14. Yorumlama'
    process_comp = data.get('process_component', '').strip() or 'Verileri ayrıştırma, parçalar arası mantıksal ilişki kurma, çıkarım yapma ve anlamlandırma'

    system_instruction = (
        'Sen; eğitim teknolojileri mimarı, TYMM bağlam temelli ölçme-değerlendirme uzmanı, '
        'Muallimin Manevi Rehberi (MMR) felsefecisi ve profesyonel eğitim materyali tasarımcısısın.\n\n'
        'Herhangi bir sınıf, ders ve ünite (Örn: 5. sınıf Fen Bilimleri → Madde ve Doğası) seçildiğinde '
        'otomatik olarak TYMM + MMR mantığıyla çalışan standart bir üretim sistemi oluşturmak üzere '
        'aşağıdaki 5 MOTORLU MİMARİYİ harfiyen uygula:\n\n'
        '1. PEDAGOJİ MOTORU:\n'
        '   TYMM → Öğrenme çıktısı + beceri + süreç bileşeni.\n'
        '   Hedeflenen bilişsel beceriyi ve süreç bileşenlerini pedagojik omurga olarak kur.\n\n'
        '2. BAĞLAM MOTORU:\n'
        '   Gerçek hayat → Problem → Veri → Durum.\n'
        '   Öğrencinin dünyasından, yaşantısal, merak uyandıran bir gerçek hayat problemi ve buna dayalı bilimsel veri seti kurgula.\n\n'
        '3. AKIL YÜRÜTME MOTORU:\n'
        '   Yorumlama → İlişkilendirme → Çıkarım → Gerekçelendirme → Değerlendirme.\n'
        '   Basitten karmaşığa giden analitik sorularla öğrencinin veriyi yorumlamasını, sebep-sonuç kurmasını ve gerekçeli karar vermesini sağla.\n\n'
        '4. MMR MANA MOTORU:\n'
        '   Bilgi → Hayret → Hikmet/Kudret → Tefekkür → Değer → Sorumluluk.\n'
        '   Bilimsel bilginin içindeki düzeni, hassas ölçüyü (mizan) ve yaratılış hikmetini hissettir. '
        '   Öğrenciye hazır dogma dayatmadan hayret ve tefekküre sevk et; ahlaki değer ve eyleme bağla.\n\n'
        '5. TASARIM MOTORU:\n'
        '   A4 → Görsel hiyerarşi → Cevap alanları → Baskı → Dijital kullanım.\n'
        '   Öğrencinin rahatça elle yazabileceği geniş çizgili [ ......................................... ] alanları bırak.\n\n'
        'ÇALIŞMA KAĞIDI KESİN YAPI KURALI:\n'
        'Çalışma kâğıdında SADECE VE SADECE aşağıdaki 7 başlık altında bölümler oluşturacaksın:\n'
        '1. BAĞLAM\n'
        '2. KANIT / VERİ / MATERYAL\n'
        '3. TYMM BECERİ GÖREVİ\n'
        '4. AKIL YÜRÜTME\n'
        '5. MMR PENCERESİ ve tefekkür\n'
        '6. Değer ve Hayata Yansıtma\n'
        '7. MÜZAKERE\n\n'
        'KESİNLİKLE BU 7 BAŞLIK DIŞINDA HİÇBİR BÖLÜM VEYA BAŞLIĞI ÇALIŞMA KAĞIDINA KOYMAYACAKSIN.\n'
        'Özellikle BAĞLAM ve AKIL YÜRÜTME kısımlarında doğrudan MMR (Muallimin Manevi Rehberi) mantığı ve diliyle konuş.'
    )

    user_prompt = (
        f'Aşağıda verilen eğitim bilgilerini 5 MOTORLU MİMARİDEN (Pedagoji Motoru → Bağlam Motoru → '
        f'Akıl Yürütme Motoru → MMR Mana Motoru → Tasarım Motoru) geçirerek ÇALIŞMA KAĞIDINDA SADECE VE SADECE '
        f'şu 7 bölümden oluşan standart materyali üret:\n\n'
        f'PEDAGOJİK GİRDİLER:\n'
        f'- Sınıf Seviyesi: {data.get("grade")}\n'
        f'- Ders: {data.get("subject")}\n'
        f'- Öğrenme Alanı / Ünite: {data.get("learning_area")}\n'
        f'- Konu: {data.get("topic")}\n'
        f'- Öğrenme Çıktısı (Kazanım): {data.get("learning_outcome")}\n'
        f'- Hedeflenen Beceri (TYMM): {skill}\n'
        f'- Süreç Bileşeni: {process_comp}\n'
        f'- Manevi Öğrenme Çıktısı (MMR): {data.get("manevi_outcome", "")}\n\n'
        f'ÇALIŞMA KAĞIDINDA SADECE AŞAĞIDAKİ 7 BAŞLIK OLACAKTIR (BAŞKA HİÇBİR BAŞLIK KOYMA):\n\n'
        f'### BAĞLAM\n'
        f'[Bağlam Motoru & MMR Mana Motoru: Gerçek hayat → problem → durum. Tabiat ve varlıktaki harikulade düzeni, ölçüyü, '
        f'hassas dengeyi ve ahengi hissettiren sahici, yaşantısal gerçek yaşam senaryosu].\n\n'
        f'### KANIT / VERİ / MATERYAL\n'
        f'[Bağlam Motoru: Problem setine ait anlaşılır bilimsel veri tablosu, grafik, deney sonucu, ölçüm değerleri veya model].\n\n'
        f'### TYMM BECERİ GÖREVİ\n'
        f'[Pedagoji Motoru: Hedeflenen bilişsel beceriyi (Karşılaştır, sınıflandır, ilişkilendir, çıkarım yap, modelle) doğrudan işleten açık görev yönergesi].\n\n'
        f'### AKIL YÜRÜTME\n'
        f'[Akıl Yürütme Motoru & MMR Mana Motoru: Yorumlama → ilişkilendirme → çıkarım → gerekçelendirme → değerlendirme. '
        f'Verilerdeki nizamı, hassas ölçüyü (mizan) ve sebep-sonuç hikmetini basitten karmaşığa analiz ettiren düşünce soruları]:\n'
        f'1. Verideki Düzeni ve İlişkiyi Fark Etme: ...\n'
        f'2. Çıkarım ve Muhakeme: ...\n'
        f'3. Karar ve Gerekçelendirme: ...\n'
        f'Çözüm ve Kararım: [ ............................................................................ ]\n\n'
        f'### MMR PENCERESİ ve tefekkür\n'
        f'[MMR Mana Motoru: Bilgi → hayret → hikmet/kudret → tefekkür]:\n'
        f'- Gözlem ve Hayretim: [ ............................................................................ ]\n'
        f'- Hikmet ve Anlamlandırmam: [ ............................................................................ ]\n'
        f'- İnsani Sorumluluğum: [ ............................................................................ ]\n\n'
        f'### Değer ve Hayata Yansıtma\n'
        f'[MMR Mana Motoru: Değer → sorumluluk → eylem]:\n'
        f'- Bu konudan fark ettiğim erdem / değer: [ ........................................................ ]\n'
        f'- Günlük hayatımda uygulayacağım somut davranış ve eylem taahhüdüm: [ ........................................................ ]\n\n'
        f'### MÜZAKERE\n'
        f'[Akıl Yürütme Motoru & MMR Mana Motoru: Sınıf içi akran diyaloğu için gerekçeli, açık uçlu düşünce soruları].\n'
        f'Görüşüm ve Savunmam: [ ............................................................................ ]\n\n'
        f'TASARIM MOTORU KURALLARI:\n'
        f'- KESİNLİKLE bu 7 başlık dışında hiçbir başlık oluşturma.\n'
        f'- Cevap alanlarında öğrencinin elle rahatça yazabilmesi için bolca [ ............................................................................ ] bırak.'
    )

    return system_instruction, user_prompt

def build_test_prompt(data: dict) -> tuple[str, str]:
    system_instruction = (
        'Sen; Türkiye Yüzyılı Maarif Modeli (TYMM), ölçme-değerlendirme ilkeleri, '
        'Erdem-Değer-Eylem (EDE) yaklaşımı ve Muallimin Manevi Rehberi (MMR) konusunda uzman '
        'bir ölçme-değerlendirme uzmanı ve deneyimli bir öğretmensin.\n\n'
        'Temel Görevin:\n'
        'Öğrencinin verilen öğrenme çıktısına ulaşma düzeyini geçerli, güvenilir ve nesnel biçimde '
        'ölçen bir YAPRAK TEST hazırlamaktır.\n\n'
        'Bilişsel Düzeyler:\n'
        'Hatırlama → Anlama → Uygulama → Analiz → Değerlendirme basamaklarını dengeli dağıt.\n\n'
        'KESİN ÖLÇME KURALLARI:\n'
        '- Yalnızca verilen öğrenme çıktısını ölç. Yeni kazanım ekleme.\n'
        '- Her soru 4 seçenekli (A, B, C, D) olmalıdır.\n'
        '- Tek ve tartışmasız bir doğru cevap bulunmalıdır.\n'
        '- Çeldiriciler bilimsel kavram yanılgılarına dayalı ve güçlü olmalıdır.\n'
        '- Hepsi veya Hiçbiri gibi seçenekleri asla kullanma.\n'
        '- Doğru cevabın uzunluğu diğer şıklardan belirgin biçimde farklı olmamalıdır.\n'
        '- Uygun sorularda günlük yaşam bağlamları, modeller veya tablolar kullan.\n'
        '- MMR boyutu için inancı doğrudan ölçen sorular sorma. Nizam, ölçü ve hikmet boyutunu akıl yürütmeye dayalı sor.\n'
        '- Testin en sonunda öğretmen için CEVAP ANAHTARI VE BİLİŞSEL DÜZEY MATRİSİ tablosu ekle.'
    )

    count = data.get('question_count', 8)
    difficulty = data.get('difficulty', 'Orta')

    user_prompt = (
        f'Aşağıda verilen eğitim bilgilerini temel alarak eksiksiz bir YAPRAK TEST hazırla:\n\n'
        f'Sınıf Seviyesi: {data.get("grade")}\n'
        f'Ders: {data.get("subject")}\n'
        f'Öğrenme Alanı: {data.get("learning_area")}\n'
        f'Konu: {data.get("topic")}\n'
        f'Öğrenme Çıktısı: {data.get("learning_outcome")}\n'
        f'Manevi Öğrenme Çıktısı: {data.get("manevi_outcome", "")}\n'
        f'İçerik Türü: Yaprak Test\n'
        f'Soru Sayısı: {count}\n'
        f'Zorluk Seviyesi: {difficulty}\n\n'
        f'Başlık Formatı: [{data.get("subject").upper()}] - [{data.get("grade").upper()}] YAPRAK TEST\n'
        f'Öğrenme Alanı, Konu, Öğrenme Çıktısı, hemen altına Manevi Öğrenme Çıktısı, Ad Soyad, Sınıf/No, Tarih, Puan kutusu ekle.\n\n'
        f'Soruları kolaydan zora sırala ve bilişsel basamakları dengeli dağıt.\n'
        f'Her soru için 4 seçenek (A, B, C, D) kullan.\n\n'
        f'SORU DÜZENİ VE FORMATI:\n'
        f'Her soruyu şu net yapıda oluştur:\n'
        f'### Soru 1\n'
        f'[Bilişsel Düzey: Hatırlama | Puan: 16.6]\n'
        f'Soru kökü ve varsa öncüller (I, II, III).\n'
        f'A) Seçenek metni\n'
        f'B) Seçenek metni [✔]\n'
        f'C) Seçenek metni\n'
        f'D) Seçenek metni\n\n'
        f'En sonda öğretmen için:\n'
        f'### CEVAP ANAHTARI VE BİLİŞSEL DÜZEY MATRİSİ\n'
        f'| Soru | Doğru Cevap | Puan | Ölçülen Süreç Bileşeni | Bilişsel Düzey |\n\n'
        f'BİÇİMLENDİRME VE YAZI KARAKTERİ KURALLARI:\n'
        f'- KESİNLİKLE LaTeX veya formül kodları ($\rightarrow$, \\rightarrow, \\to vb.) KULLANMA.\n'
        f'- Ok işaretleri için doğrudan UTF-8 "→" sembolünü kullan.\n'
        f'Çıktıyı temiz, akademik ve doğrudan yazdırılabilir Markdown olarak üret.'
    )

    return system_instruction, user_prompt

def build_case_study_prompt(data: dict) -> tuple[str, str]:
    skill = data.get('skill', '').strip() or 'KB2.4. Çözümleme / KB2.14. Yorumlama / KB2.16. Karar Verme'
    process_comp = data.get('process_component', '').strip() or 'Durum analizi yapma, kanıtları değerlendirme, gerekçeli karar alma ve etik/manevi boyutu yorumlama'

    system_instruction = (
        'Sen; eğitim teknolojileri ve vaka temelli öğretim uzmanı, problem çözme ve durum analizi tasarımcısı, '
        'TYMM beceri temelli öğretim ve Muallimin Manevi Rehberi (MMR) felsefecisisin.\n\n'
        'TEMEL GÖREVİN:\n'
        'Verilen sınıf, ders, konu ve öğrenme çıktısına uygun; öğrencinin analitik düşünmesini, '
        'durumu çok boyutlu kavramasını, kanıtlara dayalı çıkarım yapmasını ve olayın ardındaki nizam, '
        'hikmet ve ahlaki erdemleri fark etmesini sağlayan yüksek standartta bir VAKA ANALİZİ MATERYALİ üretmektir.\n\n'
        'VAKA ANALİZİ KESİN 3 BÖLÜMLÜ STANDART MİMARİSİ:\n'
        'Materyal KESİNLİKLE VE SADECE aşağıdaki 3 ana başlık altında oluşturulacaktır:\n\n'
        '1. DURUMU ANALİZ ETME\n'
        '2. ÇIKARIM VE DEĞERLENDİRME\n'
        '3. HİKMET VE DEĞER\n\n'
        'BÖLÜMLERİN İÇERİK DETAYLARI:\n'
        '- 1. DURUMU ANALİZ ETME:\n'
        '  * Vaka Senaryosu: Gerçek yaşamdan, yaşantısal, öğrencinin ilgisini çeken, ders kazanımıyla doğrudan örtüşen '
        '    somut bir olay, araştırma, problem durumu veya çevre/doğa vakası.\n'
        '  * Olayın Aktörleri, Şartları ve Temel Verileri (Maddeler veya veri tablosu halinde somut olgular).\n'
        '  * Durum Analizi Soruları: "Olayda ne oldu?", "Hangi unsurlar birbiriyle etkileşim halindedir?", "Sorunun / durumun kök sebebi nedir?".\n'
        '  * Tespit ve Analiz Alanım: [ ............................................................................ ]\n\n'
        '- 2. ÇIKARIM VE DEĞERLENDİRME:\n'
        '  * Neden-Sonuç İlişkilerinin Çözümlenmesi.\n'
        '  * Kritik Karar Noktaları ve Alternatif Çözüm Yolları.\n'
        '  * Gerekçelendirme ve Çıkarım Soruları: "Hangi karar verilmelidir ve gerekçesi nedir?", "Bu sonucun doğuracağı etkiler nelerdir?".\n'
        '  * Gerekçeli Değerlendirmem ve Nihai Kararım: [ ............................................................................ ]\n\n'
        '- 3. HİKMET VE DEĞER:\n'
        '  * MMR Boyutu: Olayın ve tabiatın arkasındaki şaşmaz düzen, ölçü (mizan), ahenk ve hikmet.\n'
        '  * Ahlaki Erdem ve Değer Keşfi: Adalet, emanet, dürüstlük, merhamet, sorumluluk veya ölçülülük.\n'
        '  * İnsani Sorumluluk ve Hayata Yansıtma: Öğrencinin bu vakadan kendi hayatına ve toplumuna çıkaracağı somut tutum.\n'
        '  * İnsani Sorumluluğum ve Eylem Taahhüdüm: [ ............................................................................ ]\n'
        '  * Akran Müzakeresi Sorusu: Sınıfta tartışılacak açık uçlu düşünce sorusu.\n\n'
        'TASARIM KURALLARI:\n'
        '- Öğrencinin elle rahatça yazabileceği geniş [ ......................................... ] çizgili alanları cömertçe bırak.\n'
        '- KESİNLİKLE bu 3 başlık dışında harici bölüm ekleme.'
    )

    user_prompt = (
        f'Aşağıda verilen eğitim bilgilerine göre eksiksiz, modern ve yazdırılabilir bir VAKA ANALİZİ MATERYALİ hazırla:\n\n'
        f'Sınıf Seviyesi: {data.get("grade")}\n'
        f'Ders: {data.get("subject")}\n'
        f'Öğrenme Alanı / Ünite: {data.get("learning_area")}\n'
        f'Konu: {data.get("topic")}\n'
        f'Öğrenme Çıktısı (Kazanım): {data.get("learning_outcome")}\n'
        f'Hedeflenen Beceri (TYMM): {skill}\n'
        f'Süreç Bileşeni: {process_comp}\n'
        f'Manevi Öğrenme Çıktısı (MMR): {data.get("manevi_outcome", "")}\n\n'
        f'MATERYALDE SADECE VE SADECE ŞU 3 BAŞLIK OLACAKTIR:\n\n'
        f'### DURUMU ANALİZ ETME\n'
        f'[Gerçek yaşam vaka senaryosu, aktörler, somut veriler, şartlar ve durum tespiti soruları]:\n'
        f'Tespit ve Analiz Alanım: [ ............................................................................ ]\n\n'
        f'### ÇIKARIM VE DEĞERLENDİRME\n'
        f'[Neden-sonuç ilişkileri, alternatif kararlar, kanıtlara dayalı analiz, gerekçelendirme soruları]:\n'
        f'Gerekçeli Değerlendirmem ve Nihai Kararım: [ ............................................................................ ]\n\n'
        f'### HİKMET VE DEĞER\n'
        f'[Vakadaki ölçü, nizam ve ilahi hikmet, ahlaki erdem keşfi, insani sorumluluk ve müzakere]:\n'
        f'- Vakadan Fark Ettiğim Değer ve Hikmet: [ ............................................................................ ]\n'
        f'- Hayata Yansıtacağım İnsani Sorumluluk: [ ............................................................................ ]\n'
        f'- Müzakere Sorusu ve Görüşüm: [ ............................................................................ ]'
    )

    return system_instruction, user_prompt

def build_mind_map_prompt(data: dict) -> tuple[str, str]:
    skill = data.get('skill', '').strip() or 'KB2.3. Örgütleme / KB2.4. Çözümleme / KB2.15. Yapılandırma'
    process_comp = data.get('process_component', '').strip() or 'Kavramları yapılandırma, örüntüleri ve hiyerarşik bağları kurma, bütüncül anlam çıkarma'

    system_instruction = (
        'Sen; görsel düşünme ve kavramsal zihin haritalama mimarı, bilişsel örgütleme uzmanı, '
        'TYMM bütüncül modelleme ve Muallimin Manevi Rehberi (MMR) şuur tasarımcısısın.\n\n'
        'TEMEL GÖREVİN:\n'
        'Verilen sınıf, ders, konu ve öğrenme çıktısına ait kavram ağını; parçadan bütüne, somuttan soyuta, '
        'nedenden sonuca sistematik olarak örgütleyen ve kâinattaki muazzam nizamla (bütünlük) birleştiren '
        'standart bir ZİHİN HARİTASI MATERYALİ üretmektir.\n\n'
        'ZİHİN HARİTASI KESİN 3 BÖLÜMLÜ STANDART MİMARİSİ:\n'
        'Materyal KESİNLİKLE VE SADECE aşağıdaki 3 ana başlık altında yapılandırılacaktır:\n\n'
        '1. ÖRGÜTLEME\n'
        '2. YAPILANDIRMA\n'
        '3. BÜTÜNLÜK\n\n'
        'BÖLÜMLERİN İÇERİK DETAYLARI:\n'
        '- 1. ÖRGÜTLEME:\n'
        '  * Merkezi Kavram ve Odak Tema: Konunun kalbindeki çekirdek ilke.\n'
        '  * Ana Dallar (Birinci Düzey Kategoriler): Konuyu oluşturan 3-4 temel sütun / kavram öbeği.\n'
        '  * Her ana dalın temel tanımı, kapsadığı alan ve ayırt edici nitelikleri.\n'
        '  * Öğrencinin Örgütleme Çalışması: [ ............................................................................ ]\n\n'
        '- 2. YAPILANDIRMA:\n'
        '  * İkinci ve Üçüncü Düzey Alt Dallar: Alt kavramlar, özellikler, formüller, süreçler ve mekanizmalar.\n'
        '  * Kavramlar Arası İlişki ve Akış Şeması (Hiyerarşik, girintili ve temiz sembollerle görselleştirilmiş kavram ağı haritası):\n'
        '    Örnek yapı:\n'
        '    └── [Merkez Kavram]\n'
        '        ├── [Ana Dal 1] ──> [Alt Kavram A] ──> [Özellik/Süreç]\n'
        '        ├── [Ana Dal 2] ──> [Alt Kavram B] ──> [İşleyiş/Denge]\n'
        '        └── [Ana Dal 3] ──> [Alt Kavram C] ──> [Hikmet/Mizan]\n'
        '  * Yapılandırma ve Çözümleme Sorusu: "Bu kavramlar birbirine nasıl bağlanır?", "Aralarındaki denge nasıl sağlanır?".\n'
        '  * Öğrencinin Kavramsal Çıkarım Alanı: [ ............................................................................ ]\n\n'
        '- 3. BÜTÜNLÜK:\n'
        '  * Parçadan Bütüne Nizam: Bütün bu kavramların ve mekanizmaların kâinattaki kusursuz ahenk ve birlik (vahdet) ile bağı.\n'
        '  * MMR Tefekkür ve Şuur Sentezi: Ayrışan parçaların arkasındaki tek ve eşsiz Yaratıcı sanatı, hikmeti ve ölçüsü.\n'
        '  * İnsani Değer ve Anlamlandırma: Bu bütünsel bakış açısının öğrencinin düşünce dünyasına ve ahlakına kattığı değer.\n'
        '  * Bütünsel Anlamlandırma ve Tefekkür Notum: [ ............................................................................ ]\n\n'
        'TASARIM KURALLARI:\n'
        '- Okunabilir, estetik, girintili hiyerarşik Markdown ve şematik ağaç düzeni kullan.\n'
        '- Öğrencinin elle doldurabileceği çizgili alanlar [ ......................................... ] bırak.\n'
        '- KESİNLİKLE bu 3 başlık dışında harici bölüm ekleme.'
    )

    user_prompt = (
        f'Aşağıda verilen eğitim parametrelerine göre eksiksiz, estetik ve öğrenci dostu bir ZİHİN HARİTASI MATERYALİ üret:\n\n'
        f'Sınıf Seviyesi: {data.get("grade")}\n'
        f'Ders: {data.get("subject")}\n'
        f'Öğrenme Alanı / Ünite: {data.get("learning_area")}\n'
        f'Konu: {data.get("topic")}\n'
        f'Öğrenme Çıktısı (Kazanım): {data.get("learning_outcome")}\n'
        f'Hedeflenen Beceri (TYMM): {skill}\n'
        f'Süreç Bileşeni: {process_comp}\n'
        f'Manevi Öğrenme Çıktısı (MMR): {data.get("manevi_outcome", "")}\n\n'
        f'MATERYALDE KESİNLİKLE VE SADECE ŞU 3 ANA BAŞLIK OLACAKTIR:\n\n'
        f'### ÖRGÜTLEME\n'
        f'[Merkezi kavram, 3-4 temel ana dal, kavramsal dayanaklar ve sınıflandırma şeması]:\n'
        f'Kavramsal Odak ve Örgütleme Notum: [ ............................................................................ ]\n\n'
        f'### YAPILANDIRMA\n'
        f'[Alt dallar, kavramlar arası ilişkiler, süreçler ve sebep-sonuç ağları]:\n'
        f'```text\n'
        f'└── [{data.get("topic")}]\n'
        f'    ├── [1. Ana Dal] ──> [Alt Kavram/Özellik] ──> [Süreç/Sonuç]\n'
        f'    ├── [2. Ana Dal] ──> [Alt Kavram/Özellik] ──> [Süreç/Sonuç]\n'
        f'    └── [3. Ana Dal / Hikmet] ──> [Kusursuz Nizam] ──> [Tefekkür]\n'
        f'```\n'
        f'Kavramlar Arası İlişki ve Çözümlemem: [ ............................................................................ ]\n\n'
        f'### BÜTÜNLÜK\n'
        f'[Parçalardan bütüne kâinattaki ahenk, nizam ve birlik (vahdet); MMR tefekkür sentezi ve ahlaki şuur]:\n'
        f'- Bütünsel Nizam ve Hikmet Tefekkürüm: [ ............................................................................ ]\n'
        f'- Bu Kavram Ağından Kazandığım Değer ve Şuur: [ ............................................................................ ]'
    )

    return system_instruction, user_prompt

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/calisma_kagidi_bankasi.html')
def bankasi_view():
    base_dir = os.path.dirname(os.path.abspath(__file__))
    file_path = os.path.join(base_dir, 'calisma_kagidi_bankasi.html')
    if os.path.exists(file_path):
        return send_file(file_path)
    return "Dosya bulunamadı", 404

@app.route('/materyal_goruntuleyici.html')
def materyal_view():
    base_dir = os.path.dirname(os.path.abspath(__file__))
    file_path = os.path.join(base_dir, 'materyal_goruntuleyici.html')
    if os.path.exists(file_path):
        return send_file(file_path)
    return "Dosya bulunamadı", 404

@app.route('/calisma_kagidi_bankasi/<path:filename>')
def download_bank_docx(filename):
    base_dir = os.path.dirname(os.path.abspath(__file__))
    file_path = os.path.join(base_dir, 'calisma_kagidi_bankasi', filename)
    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=filename.endswith('.docx'))
    return "Dosya bulunamadı", 404

@app.route('/api/health')
def health():
    key = get_api_key()
    has_key = bool(key and len(key) > 10)
    return jsonify({
        'status': 'ok',
        'has_api_key': has_key
    })

def load_mmr_unite_plani_units():
    docx_path = os.path.join(os.path.dirname(__file__), 'Muallimin_Manevi_Rehberi_1_Unite_Plani.docx')
    if os.path.exists(docx_path):
        try:
            import docx
            doc = docx.Document(docx_path)
            table = doc.tables[0]
            units = []
            
            week_titles = [
                "1. Hafta: Güneş'in Yapısı ve Canlılığın Devamı (F.M.5.1.1.1)",
                "2. Hafta: Güneş'in Dönme Hareketi ve Gezegenlere Rahmet Olması (F.M.5.1.1.1)",
                "3. Hafta: Ay'ın Özellikleri ve Hareketlerindeki Mükemmel Düzen (F.M.5.1.2.1)",
                "4. Hafta: Ay'ın Evreleri ve Zaman Ölçüsündeki İlahi İntizam (F.M.5.1.2.2)",
                "5. Hafta: Güneş, Dünya ve Ay'ın Muazzam Uyumu ve Vahdet (F.M.5.1.3.1)"
            ]
            
            week_topics = [
                "Güneş'in Yapısı ve Canlılığın Devamındaki Mükemmel Yaratılış",
                "Güneş'in Dönme Hareketi ve Dünyamıza Rahmet Olması",
                "Ay'ın Özellikleri, Dönme ve Dolanma Hareketlerindeki Mükemmel Düzen",
                "Ay'ın Evreleri ve Zaman Ölçüsündeki İlahi İntizam",
                "Güneş, Dünya ve Ay'ın Birbirine Göre Hareketleri ve Eşsiz Uyumu"
            ]

            for idx, row in enumerate(table.rows[1:]):
                cells = [c.text.replace('\r', '').strip() for c in row.cells]
                if len(cells) >= 5:
                    outcome_raw = cells[1].replace('\n', ' ')
                    proc_comp = cells[2].replace('\n', ' ')
                    degerler = cells[3].replace('\n', ' ')
                    manevi_outcome = cells[4].replace('\n', ' ')
                    
                    unit = {
                        'title': week_titles[idx] if idx < len(week_titles) else f"{cells[0]}: {cells[1][:40]}...",
                        'grade': '5. Sınıf',
                        'subject': 'Fen Bilimleri',
                        'learning_area': '1. Ünite: Güneş, Dünya ve Ay (Dünya ve Evren)',
                        'topic': week_topics[idx] if idx < len(week_topics) else cells[1].split('\n')[-1],
                        'learning_outcome': outcome_raw,
                        'skill': f'TYMM Becerisi / Değerler: {degerler}',
                        'process_component': proc_comp,
                        'manevi_outcome': manevi_outcome
                    }
                    units.append(unit)
            if units:
                return units
        except Exception as e:
            print("Docx parse error, fallback to hardcoded MMR plan:", e)

    return [
        {
            'title': "1. Hafta: Güneş'in Yapısı ve Canlılığın Devamı (F.M.5.1.1.1)",
            'grade': '5. Sınıf',
            'subject': 'Fen Bilimleri',
            'learning_area': '1. Ünite: Güneş, Dünya ve Ay (Dünya ve Evren)',
            'topic': "Güneş'in Yapısı ve Canlılığın Devamındaki Mükemmel Yaratılış",
            'learning_outcome': "F.M.5.1.1.1 Güneş'in yapısı ve dönme hareketi ile ilgili bilgileri toplayabilme.",
            'skill': 'Bilgi Toplama / Sorumluluk, Hikmet, Tefekkür, Şükür',
            'process_component': 'a) Bilgiye ulaşmak için araçları belirler. b) Belirlediği araçları kullanarak bilgileri bulur. c) Doğrular. ç) Kaydeder.',
            'manevi_outcome': "Güneş'in yapısı ve hareketleri bakımından canlılığın devamına katkısındaki mükemmel yaratılışını fark edebilme."
        },
        {
            'title': "2. Hafta: Güneş'in Dönme Hareketi ve Gezegenlere Rahmet Olması (F.M.5.1.1.1)",
            'grade': '5. Sınıf',
            'subject': 'Fen Bilimleri',
            'learning_area': '1. Ünite: Güneş, Dünya ve Ay (Dünya ve Evren)',
            'topic': "Güneş'in Dönme Hareketi ve Dünyamıza Rahmet Olması",
            'learning_outcome': "F.M.5.1.1.1 Güneş'in yapısı ve dönme hareketi ile ilgili bilgileri toplayabilme.",
            'skill': 'Bilimsel Gözlem ve Kayıt / Merhamet, İntizam, Tevhid, İbret',
            'process_component': 'a) Bilgiye ulaşmak için araçları belirler. b) Bilgileri bulur. c) Doğrular. ç) Ulaşılan bilgileri kaydeder.',
            'manevi_outcome': "Güneş'in hareketli bir varlık olarak yaratılmasının etrafındaki gezegenlere, özellikle dünyamıza, bir rahmet olduğunu kavrayabilme."
        },
        {
            'title': "3. Hafta: Ay'ın Özellikleri ve Hareketlerindeki Mükemmel Düzen (F.M.5.1.2.1)",
            'grade': '5. Sınıf',
            'subject': 'Fen Bilimleri',
            'learning_area': '1. Ünite: Güneş, Dünya ve Ay (Dünya ve Evren)',
            'topic': "Ay'ın Özellikleri, Dönme ve Dolanma Hareketlerindeki Mükemmel Düzen",
            'learning_outcome': "F.M.5.1.2.1 Ay'ın özellikleri, dönme ve dolanma hareketleri ile ilgili bilimsel çıkarım yapabilme.",
            'skill': 'Bilimsel Çıkarım / Adalet, Denge, Tevhid, Hayret ve Hayranlık',
            'process_component': "a) Ay'ın özellikleri ve hareketleri ile ilgili nitelikleri tanımlar. b) Topladığı verileri kaydeder. c) Verileri değerlendirir.",
            'manevi_outcome': "Ay'ın hareketlerindeki mükemmel düzeni fark ederek, bu düzenin sonsuz güç sahibi Allah tarafından sağlandığı hakkında çıkarım yapabilme."
        },
        {
            'title': "4. Hafta: Ay'ın Evreleri ve Zaman Ölçüsündeki İlahi İntizam (F.M.5.1.2.2)",
            'grade': '5. Sınıf',
            'subject': 'Fen Bilimleri',
            'learning_area': '1. Ünite: Güneş, Dünya ve Ay (Dünya ve Evren)',
            'topic': "Ay'ın Evreleri ve Zaman Ölçüsündeki İlahi İntizam",
            'learning_outcome': "F.M.5.1.2.2 Ay'ın evrelerini temsil eden bilimsel model oluşturabilme.",
            'skill': 'Bilimsel Modelleme / Zaman Bilinci, Şükür, Hikmet, İntizam',
            'process_component': "a) Ay'ın evrelerini temsil eden bir model önerir. b) Modelini yeni kanıtlara bağlı olarak geliştirir.",
            'manevi_outcome': "Ay'ın hareketlerindeki ince hesap ve düzenden yola çıkarak Cenab-ı Hakkın her şeye gücünün yettiğini anlayabilme."
        },
        {
            'title': "5. Hafta: Güneş, Dünya ve Ay'ın Muazzam Uyumu ve Vahdet (F.M.5.1.3.1)",
            'grade': '5. Sınıf',
            'subject': 'Fen Bilimleri',
            'learning_area': '1. Ünite: Güneş, Dünya ve Ay (Dünya ve Evren)',
            'topic': "Güneş, Dünya ve Ay'ın Birbirine Göre Hareketleri ve Eşsiz Uyumu",
            'learning_outcome': "F.M.5.1.3.1 Güneş, Dünya ve Ay'ın birbirlerine göre hareketlerini ve hacimsel büyüklüklerini temsil eden bilimsel model oluşturabilme.",
            'skill': 'Sistemik Modelleme / Uyum, Birlik (Vahdet), İntizam, Sanat',
            'process_component': "a) Güneş, Dünya ve Ay'ın birbirlerine göre hareketlerini ve büyüklüklerini temsil eden bir model önerir. b) Modelini geliştirir.",
            'manevi_outcome': "Güneş, Dünya ve Ay arasındaki harika uyum ve hareketlerin Allah'ın varlığına işaret ettiği çıkarımında bulunabilme."
        }
    ]

@app.route('/api/sample-units')
def sample_units():
    units = load_mmr_unite_plani_units()
    return jsonify({'units': units})

@app.route('/api/generate', methods=['POST'])
def generate():
    try:
        data = request.get_json() or {}
        grade = data.get('grade', '').strip()
        subject = data.get('subject', '').strip()
        learning_area = data.get('learning_area', '').strip()
        topic = data.get('topic', '').strip()
        learning_outcome = data.get('learning_outcome', '').strip()
        content_type = data.get('content_type', 'worksheet').strip()

        if not all([grade, subject, learning_area, topic, learning_outcome]):
            return jsonify({
                'success': False,
                'error': 'Lütfen tüm zorunlu alanları (Sınıf, Ders, Öğrenme Alanı, Konu, Öğrenme Çıktısı) doldurunuz.'
            }), 400

        if content_type == 'worksheet':
            sys_inst, user_prompt = build_worksheet_prompt(data)
        elif content_type == 'test':
            sys_inst, user_prompt = build_test_prompt(data)
        elif content_type == 'case_study':
            sys_inst, user_prompt = build_case_study_prompt(data)
        elif content_type == 'mind_map':
            sys_inst, user_prompt = build_mind_map_prompt(data)
        else:
            sys_inst, user_prompt = build_worksheet_prompt(data)

        client_key = request.headers.get('X-Gemini-Key') or data.get('api_key')
        content = call_gemini_api(sys_inst, user_prompt, client_key)
        
        return jsonify({
            'success': True,
            'content_type': content_type,
            'content': content
        })

    except ValueError as e:
        logger.error(f'Validation/Key error: {str(e)}')
        return jsonify({'success': False, 'error': str(e)}), 400
    except Exception as e:
        logger.error(f'Generation error: {str(e)}')
        return jsonify({
            'success': False,
            'error': 'İçerik oluşturulurken bir sorun oluştu. Lütfen bilgileri kontrol ederek tekrar deneyiniz.'
        }), 500

@app.route('/api/export-docx', methods=['POST'])
def export_docx():
    try:
        from docx import Document
        from docx.shared import Pt, Inches

        data = request.get_json() or {}
        text = data.get('content', '')
        title = data.get('title', 'Materyal')

        doc = Document()
        for section in doc.sections:
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

        lines = text.split('\n')
        for line in lines:
            line_str = line.strip()
            if not line_str:
                doc.add_paragraph()
                continue
            
            if line_str.startswith('# '):
                h = doc.add_heading(line_str[2:], level=1)
                h.paragraph_format.space_before = Pt(12)
                h.paragraph_format.space_after = Pt(6)
            elif line_str.startswith('## '):
                h = doc.add_heading(line_str[3:], level=2)
                h.paragraph_format.space_before = Pt(10)
                h.paragraph_format.space_after = Pt(4)
            elif line_str.startswith('### '):
                h = doc.add_heading(line_str[4:], level=3)
                h.paragraph_format.space_before = Pt(8)
                h.paragraph_format.space_after = Pt(2)
            elif line_str.startswith('#### '):
                h = doc.add_heading(line_str[5:], level=4)
                h.paragraph_format.space_before = Pt(6)
                h.paragraph_format.space_after = Pt(2)
            elif line_str.startswith('* ') or line_str.startswith('- '):
                p = doc.add_paragraph(line_str[2:], style='List Bullet')
                p.paragraph_format.space_after = Pt(2)
            elif line_str.startswith('> '):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                run = p.add_run(line_str[2:])
                run.italic = True
            elif line_str.startswith('|') and line_str.endswith('|'):
                p = doc.add_paragraph(line_str)
                p.paragraph_format.space_after = Pt(2)
            else:
                p = doc.add_paragraph(line_str)
                p.paragraph_format.space_after = Pt(4)

        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        
        filename = f"{title.replace(' ', '_')}.docx"
        return send_file(
            bio,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        logger.error(f'DOCX export error: {str(e)}')
        return jsonify({'success': False, 'error': 'Word dosyası oluşturulamadı.'}), 500

if __name__ == '__main__':
    logger.info('MMR Yaprak Test & Çalışma Kağıdı Üretim Sunucusu Başlatılıyor: http://127.0.0.1:5000')
    app.run(host='127.0.0.1', port=5000, debug=False)
