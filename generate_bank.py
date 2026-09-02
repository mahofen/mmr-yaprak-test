import os
import json
import re
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

load_dotenv()
from app import call_gemini_api, build_worksheet_prompt

UNITS = [
    {
        'week_num': 1,
        'title': "1. Hafta: Güneş'in Yapısı ve Dönme Hareketi (Soba-Lamba Analojisi)",
        'grade': '5. Sınıf',
        'subject': 'Fen Bilimleri',
        'learning_area': 'Dünya ve Evren',
        'topic': "Güneş'in Yapısı ve Dönme Hareketi",
        'learning_outcome': "F.M.5.1.1.1. Güneş'in yapısı ve dönme hareketi ile ilgili bilgileri toplayabilme.",
        'manevi_outcome': "Güneş'in yapısı ve hareketleri bakımından canlılığın devamına katkısındaki mükemmel yaratılışını fark edebilme.",
        'slug': "01_Hafta_1_Gunesin_Yapisi"
    },
    {
        'week_num': 2,
        'title': "2. Hafta: Güneş'in Dönme Hareketi ve Rahmet Boyutu",
        'grade': '5. Sınıf',
        'subject': 'Fen Bilimleri',
        'learning_area': 'Dünya ve Evren',
        'topic': "Güneş'in Dönme Hareketinin Canlılığa Etkisi",
        'learning_outcome': "F.M.5.1.1.1. Güneş'in yapısı ve dönme hareketi ile ilgili bilgileri toplayabilme ve canlılığa katkısını değerlendirebilme.",
        'manevi_outcome': "Güneş'in hareketli bir varlık olarak yaratılmasının etrafındaki gezegenlere, özellikle dünyamıza, bir rahmet olduğunu kavrayabilme.",
        'slug': "02_Hafta_2_Gunesin_Donme_Hareketi_Rahmet"
    },
    {
        'week_num': 3,
        'title': "3. Hafta: Ay'ın Özellikleri ve Hareketleri (Lisan-ı Hal)",
        'grade': '5. Sınıf',
        'subject': 'Fen Bilimleri',
        'learning_area': 'Dünya ve Evren',
        'topic': "Ay'ın Özellikleri, Dönme ve Dolanma Hareketleri",
        'learning_outcome': "F.M.5.1.2.1. Ay'ın özellikleri, dönme ve dolanma hareketleri ile ilgili bilimsel çıkarım yapabilme.",
        'manevi_outcome': "Ay'ın hareketlerindeki mükemmel düzeni fark ederek, bu düzenin sonsuz güç sahibi Allah tarafından sağlandığı hakkında çıkarım yapabilme.",
        'slug': "03_Hafta_3_Ayin_Ozellikleri_ve_Hareketleri"
    },
    {
        'week_num': 4,
        'title': "4. Hafta: Ay'ın Evreleri (Gökyüzündeki İlahi Takvim)",
        'grade': '5. Sınıf',
        'subject': 'Fen Bilimleri',
        'learning_area': 'Dünya ve Evren',
        'topic': "Ay'ın Evreleri ve Zaman Hesaplama",
        'learning_outcome': "F.M.5.1.2.2. Ay'ın evrelerini temsil eden bilimsel model oluşturabilme.",
        'manevi_outcome': "Ay'ın hareketlerindeki ince hesap ve düzenden yola çıkarak Cenab-ı Hakkın her şeye gücünün yettiğini anlayabilme.",
        'slug': "04_Hafta_4_Ayin_Evreleri_Ilahi_Takvim"
    },
    {
        'week_num': 5,
        'title': "5. Hafta: Güneş, Dünya ve Ay'ın Muazzam Uyumu (Vahdet ve Sanat)",
        'grade': '5. Sınıf',
        'subject': 'Fen Bilimleri',
        'learning_area': 'Dünya ve Evren',
        'topic': "Güneş, Dünya ve Ay'ın Birbirlerine Göre Hareketleri ve Büyüklük Modelleri",
        'learning_outcome': "F.M.5.1.3.1. Güneş, Dünya ve Ay'ın birbirlerine göre hareketlerini ve hacimsel büyüklüklerini temsil eden bilimsel model oluşturabilme.",
        'manevi_outcome': "Güneş, Dünya ve Ay arasındaki harika uyum ve hareketlerin Allah'ın varlığına işaret ettiği çıkarımında bulunabilme.",
        'slug': "05_Hafta_5_Gunes_Dunya_Ay_Muazzam_Uyum"
    }
]

def sanitize_text(text: str) -> str:
    if not text:
        return ''
    text = text.replace('\\rightarrow', '→')
    text = text.replace('$\\rightarrow$', '→')
    text = text.replace('\\to', '→')
    text = text.replace('$\\to$', '→')
    text = re.sub(r'\$([^\$]+)\$', r'\1', text)
    return text

def save_markdown_to_docx(text: str, docx_path: str, title: str):
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    header_p = doc.add_paragraph()
    run_h = header_p.add_run(f"TÜRKİYE YÜZYILI MAARİF MODELİ & MUALLİMİN MANEVİ REHBERİ (MMR)\n{title.upper()}\n")
    run_h.bold = True
    run_h.font.size = Pt(12)
    run_h.font.color.rgb = RGBColor(30, 58, 138)
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    lines = text.split('\n')
    for line in lines:
        line_str = line.strip()
        if not line_str:
            doc.add_paragraph()
            continue

        if line_str.startswith('# '):
            h = doc.add_heading(line_str[2:], level=1)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(4)
        elif line_str.startswith('## '):
            h = doc.add_heading(line_str[3:], level=2)
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(4)
        elif line_str.startswith('### '):
            h = doc.add_heading(line_str[4:], level=3)
            h.paragraph_format.space_before = Pt(8)
            h.paragraph_format.space_after = Pt(2)
        elif line_str.startswith('* ') or line_str.startswith('- '):
            p = doc.add_paragraph(line_str[2:], style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
        elif line_str.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            run = p.add_run(line_str[2:])
            run.italic = True
        elif line_str.startswith('|') and line_str.endswith('|'):
            p = doc.add_paragraph(line_str)
            p.paragraph_format.space_after = Pt(2)
        else:
            p = doc.add_paragraph(line_str)
            p.paragraph_format.space_after = Pt(4)

    doc.save(docx_path)
    print(f"[DOCX] Kaydedildi: {docx_path}")

def generate_all():
    output_dir = 'calisma_kagidi_bankasi'
    os.makedirs(output_dir, exist_ok=True)
    
    generated_worksheets = []
    
    print("=" * 60)
    print("MMR ÇALIŞMA KAĞIDI BANKASI ÜRETİM MOTORU BAŞLADI")
    print(f"{len(UNITS)} adet haftalık çalışma kağıdı üretiliyor...")
    print("=" * 60)

    for unit in UNITS:
        print(f"\n---> [{unit['week_num']}/5] {unit['title']} üretiliyor...")
        sys_inst, user_prompt = build_worksheet_prompt(unit)
        raw_text = call_gemini_api(sys_inst, user_prompt)
        clean_text = sanitize_text(raw_text)

        md_filename = os.path.join(output_dir, f"{unit['slug']}.md")
        with open(md_filename, 'w', encoding='utf-8') as f:
            f.write(clean_text)
        print(f"[MD] Kaydedildi: {md_filename} ({len(clean_text)} karakter)")

        docx_filename = os.path.join(output_dir, f"{unit['slug']}.docx")
        save_markdown_to_docx(clean_text, docx_filename, unit['title'])

        generated_worksheets.append({
            'unit': unit,
            'markdown': clean_text,
            'md_filename': f"{unit['slug']}.md",
            'docx_filename': f"{unit['slug']}.docx"
        })

    master_docx_path = os.path.join(output_dir, "CALISMA_KAGIDI_BANKASI_TUMU.docx")
    master_doc = docx.Document()
    
    p_cover = master_doc.add_paragraph()
    p_cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p_cover.add_run("TÜRKİYE YÜZYILI MAARİF MODELİ\nMUALLİMİN MANEVİ REHBERİ (MMR)\n\n")
    r1.bold = True
    r1.font.size = Pt(16)
    r1.font.color.rgb = RGBColor(15, 118, 110)

    r2 = p_cover.add_run("5. SINIF FEN BİLİMLERİ\nÇALIŞMA KAĞIDI BANKASI (5 HAFTALIK TAM SET)\n\n")
    r2.bold = True
    r2.font.size = Pt(20)
    r2.font.color.rgb = RGBColor(30, 58, 138)

    r3 = p_cover.add_run("1. Ünite: Güneş, Dünya ve Ay\nPedagojik Omurga: Bilgi → Düşünme → Anlam → Tefekkür → Şuur → Erdem → Değer → Eylem\n\n")
    r3.italic = True
    r3.font.size = Pt(11)

    master_doc.add_page_break()

    for idx, item in enumerate(generated_worksheets):
        if idx > 0:
            master_doc.add_page_break()
        
        h_unit = master_doc.add_heading(item['unit']['title'], level=1)
        h_unit.paragraph_format.space_before = Pt(14)
        h_unit.paragraph_format.space_after = Pt(8)

        lines = item['markdown'].split('\n')
        for line in lines:
            line_str = line.strip()
            if not line_str:
                master_doc.add_paragraph()
                continue

            if line_str.startswith('# '):
                h = master_doc.add_heading(line_str[2:], level=1)
            elif line_str.startswith('## '):
                h = master_doc.add_heading(line_str[3:], level=2)
            elif line_str.startswith('### '):
                h = master_doc.add_heading(line_str[4:], level=3)
            elif line_str.startswith('* ') or line_str.startswith('- '):
                master_doc.add_paragraph(line_str[2:], style='List Bullet')
            elif line_str.startswith('> '):
                p = master_doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                run = p.add_run(line_str[2:])
                run.italic = True
            elif line_str.startswith('|') and line_str.endswith('|'):
                master_doc.add_paragraph(line_str)
            else:
                master_doc.add_paragraph(line_str)

    master_doc.save(master_docx_path)
    print(f"\n[MASTER DOCX] Bütünleşik Banka Kaydedildi: {master_docx_path}")

    index_json_path = os.path.join(output_dir, "bank_data.json")
    with open(index_json_path, 'w', encoding='utf-8') as f:
        json.dump(generated_worksheets, f, ensure_ascii=False, indent=2)
    print(f"[JSON] Veri İndeksi Kaydedildi: {index_json_path}")

    print("\n" + "=" * 60)
    print("5 HAFTALIK ÇALIŞMA KAĞIDI BANKASI EKSİKSİZ OLUŞTURULDU!")
    print("=" * 60)

if __name__ == '__main__':
    generate_all()
